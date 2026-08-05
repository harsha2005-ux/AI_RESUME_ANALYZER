import os
import re
import json
from datetime import datetime

from flask import Flask, render_template, request, redirect, url_for, session, flash
from flask_sqlalchemy import SQLAlchemy
from werkzeug.security import generate_password_hash, check_password_hash
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader
from docx import Document

app = Flask(__name__)
app.config["SECRET_KEY"] = "change-this-secret-key"
app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///resume.db"
app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
app.config["UPLOAD_FOLDER"] = "uploads"
app.config["MAX_CONTENT_LENGTH"] = 10 * 1024 * 1024
os.makedirs(app.config["UPLOAD_FOLDER"], exist_ok=True)

db = SQLAlchemy(app)

SKILLS = {
    "Programming": ["python", "java", "c++", "c", "javascript", "typescript", "php"],
    "Web": ["html", "css", "react", "angular", "node.js", "flask", "django", "bootstrap"],
    "Database": ["sql", "mysql", "postgresql", "mongodb", "oracle", "excel"],
    "Data/AI": ["machine learning", "deep learning", "data science", "pandas", "numpy", "tensorflow", "pytorch", "statistics", "data visualization"],
    "Cloud/DevOps": ["aws", "azure", "gcp", "docker", "kubernetes", "git", "github", "linux"],
    "Soft Skills": ["communication", "teamwork", "leadership", "problem solving", "time management"]
}

JOBS = {
    "Software Developer": ["python", "java", "javascript", "html", "css", "sql", "git"],
    "Data Analyst": ["python", "sql", "excel", "pandas", "data visualization", "statistics"],
    "Web Developer": ["html", "css", "javascript", "react", "node.js", "sql", "git"],
    "Machine Learning Engineer": ["python", "machine learning", "pandas", "numpy", "tensorflow", "sql"],
    "Cloud Engineer": ["aws", "azure", "linux", "docker", "kubernetes", "python"]
}

COURSES = {
    "python": "Python for Everybody",
    "sql": "SQL for Data Analysis",
    "react": "React - The Complete Guide",
    "machine learning": "Machine Learning Specialization",
    "aws": "AWS Cloud Practitioner",
    "docker": "Docker & Kubernetes",
    "data science": "IBM Data Science Professional Certificate",
    "javascript": "JavaScript Algorithms and Data Structures",
    "git": "Git and GitHub Essentials",
    "excel": "Excel Skills for Business",
    "pandas": "Data Analysis with Python and Pandas",
    "statistics": "Statistics for Data Science"
}


class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    email = db.Column(db.String(120), unique=True, nullable=False)
    password = db.Column(db.String(255), nullable=False)
    role = db.Column(db.String(20), default="user")
    resumes = db.relationship("Resume", backref="user", lazy=True)


class Resume(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    filename = db.Column(db.String(255))
    extracted_text = db.Column(db.Text)
    skills = db.Column(db.Text)
    missing_skills = db.Column(db.Text)
    ats_score = db.Column(db.Integer)
    created_at = db.Column(db.DateTime, default=datetime.utcnow)
    user_id = db.Column(db.Integer, db.ForeignKey("user.id"), nullable=False)


def extract_text(path):
    if path.lower().endswith(".pdf"):
        reader = PdfReader(path)
        return "\n".join(
            (page.extract_text() or "")
            for page in reader.pages
            if (page.extract_text() or "").strip()
        )

    if path.lower().endswith(".docx"):
        document = Document(path)
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts)

    return ""


def find_skills(text):
    text_lower = text.lower()
    found = []
    for skill_list in SKILLS.values():
        for skill in skill_list:
            pattern = r"(?<!\w)" + re.escape(skill.lower()) + r"(?!\w)"
            if re.search(pattern, text_lower):
                found.append(skill)
    return sorted(set(found))


def get_best_job(skills):
    return max(JOBS, key=lambda job: len(set(skills) & set(JOBS[job])))


def build_courses(missing):
    result = [COURSES[s] for s in missing if s in COURSES][:5]
    return result or ["Improve your projects, GitHub portfolio, and interview communication."]


def build_questions(skills):
    questions = [f"Explain a project where you used {s}." for s in skills[:6]]
    questions += [
        "Tell me about yourself.",
        "Describe a challenging problem and how you solved it.",
        "Why should we hire you?"
    ]
    return questions[:10]


def analyze(text):
    found = find_skills(text)
    job = get_best_job(found)
    missing = [s for s in JOBS[job] if s not in found]
    score = 35 + len(found) * 4 + min(len(text) // 500, 15)
    if re.search(r"experience|project", text, re.I):
        score += 10
    score = min(100, score)
    return found, missing, score, job, build_courses(missing), build_questions(found)


def login_required():
    return "user_id" in session


@app.route("/")
def index():
    return render_template("index.html")


@app.route("/register", methods=["GET", "POST"])
def register():
    if request.method == "POST":
        name = request.form["name"].strip()
        email = request.form["email"].lower().strip()
        password = request.form["password"]

        if User.query.filter_by(email=email).first():
            flash("Email already registered.")
            return redirect(url_for("register"))

        db.session.add(User(
            name=name,
            email=email,
            password=generate_password_hash(password)
        ))
        db.session.commit()
        flash("Registration successful. Please log in.")
        return redirect(url_for("login"))

    return render_template("auth.html", mode="register")


@app.route("/login", methods=["GET", "POST"])
def login():
    if request.method == "POST":
        user = User.query.filter_by(
            email=request.form["email"].lower().strip()
        ).first()

        if user and check_password_hash(user.password, request.form["password"]):
            session["user_id"] = user.id
            session["role"] = user.role
            session["name"] = user.name
            return redirect(url_for("dashboard"))

        flash("Invalid email or password.")

    return render_template("auth.html", mode="login")


@app.route("/logout")
def logout():
    session.clear()
    return redirect(url_for("index"))


@app.route("/dashboard")
def dashboard():
    if not login_required():
        return redirect(url_for("login"))

    resumes = Resume.query.filter_by(
        user_id=session["user_id"]
    ).order_by(Resume.created_at.desc()).all()

    return render_template("dashboard.html", resumes=resumes)


@app.route("/analyze", methods=["POST"])
def analyze_resume():
    if not login_required():
        return redirect(url_for("login"))

    uploaded_file = request.files.get("resume")

    if not uploaded_file or not uploaded_file.filename:
        flash("Choose a PDF or DOCX file.")
        return redirect(url_for("dashboard"))

    filename = secure_filename(uploaded_file.filename)

    if not filename.lower().endswith((".pdf", ".docx")):
        flash("Only PDF and DOCX files are supported.")
        return redirect(url_for("dashboard"))

    file_path = os.path.join(
        app.config["UPLOAD_FOLDER"],
        f"{datetime.now().timestamp()}_{filename}"
    )

    try:
        uploaded_file.save(file_path)
        extracted_text = extract_text(file_path)

        if not extracted_text.strip():
            raise ValueError(
                "No readable text was found. If this is a scanned PDF, "
                "use a text-based PDF or add OCR support."
            )

        skills, missing, score, job, courses, questions = analyze(extracted_text)

        resume = Resume(
            filename=filename,
            extracted_text=extracted_text,
            skills=json.dumps(skills),
            missing_skills=json.dumps(missing),
            ats_score=score,
            user_id=session["user_id"]
        )

        db.session.add(resume)
        db.session.commit()

        return redirect(url_for("result", resume_id=resume.id))

    except Exception as error:
        db.session.rollback()
        flash(f"Could not analyze the file: {error}")
        return redirect(url_for("dashboard"))


@app.route("/result/<int:resume_id>")
def result(resume_id):
    if not login_required():
        return redirect(url_for("login"))

    resume = Resume.query.filter_by(
        id=resume_id,
        user_id=session["user_id"]
    ).first_or_404()

    skills = json.loads(resume.skills or "[]")
    missing = json.loads(resume.missing_skills or "[]")

    data = {
        "id": resume.id,
        "skills": skills,
        "missing": missing,
        "score": resume.ats_score,
        "job": get_best_job(skills),
        "courses": build_courses(missing),
        "questions": build_questions(skills),
        "text": resume.extracted_text[:4000]
    }

    return render_template("result.html", r=data)


@app.route("/history/<int:resume_id>")
def history_item(resume_id):
    return redirect(url_for("result", resume_id=resume_id))


@app.route("/admin")
def admin():
    if session.get("role") != "admin":
        flash("Admin access required.")
        return redirect(url_for("dashboard"))

    users = User.query.all()
    resumes = Resume.query.order_by(Resume.created_at.desc()).all()

    return render_template("admin.html", users=users, resumes=resumes)


@app.cli.command("init-db")
def init_db():
    db.create_all()

    if not User.query.filter_by(email="admin@resumeai.com").first():
        db.session.add(User(
            name="Admin",
            email="admin@resumeai.com",
            password=generate_password_hash("Admin@123"),
            role="admin"
        ))
        db.session.commit()

    print("Database initialized.")
    print("Admin email: admin@resumeai.com")
    print("Admin password: Admin@123")


if __name__ == "__main__":
    with app.app_context():
        db.create_all()

    app.run(debug=True)